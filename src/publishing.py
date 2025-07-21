import docx
from config import settings
from markdowntodocx.markdownconverter import convertMarkdownInFile

class DocWriter:

    def __init__(self,sett) -> None:
        self.doc = docx.Document()
        settings=settings



    def write_doc(self, book, chapter_dict, title):

        self.doc.add_heading(title, 0)

        for chapter, paragraphs_list in book.items():

            description = chapter_dict[chapter] if chapter in chapter_dict else " "
            chapter_name = "{}: {}".format(chapter.strip(), description.strip())

            self.doc.add_page_break()
            self.doc.add_heading(chapter_name, 1)

            text = "\n\n".join(paragraphs_list)
            self.doc.add_paragraph(text)

        outf = settings["outdoc"] if settings["outdoc"]!="" else "book.docx"
        self.doc.save(f"./docs/{outf}")
        convertMarkdownInFile(f"./docs/{outf}", f"./docs/{outf}")
    
    def write_structure(self, title, chapter_dict, summaries_dict, ideas_dict):
        self.doc_i = docx.Document()
        self.doc_i.add_heading(title,0)
        for chapter,ch_title in chapter_dict.items():
            summary = summaries_dict[chapter]
            ideas = ideas_dict[chapter]
            self.doc_i.add_page_break()
            self.doc_i.add_heading(f"{chapter}: {ch_title}",1)
            self.doc_i.add_paragraph(summary)
            self.doc_i.add_heading("Идеи:")
            for idea in ideas:
                self.doc_i.add_paragraph(idea, style='List Bullet')

        outf = settings["outdoc_ideas"] if settings["outdoc_ideas"]!="" else "book_ideas.docx"
        self.doc_i.save(f"./docs/{outf}")
        convertMarkdownInFile(f"./docs/{outf}", f"./docs/{outf}")

                                   



